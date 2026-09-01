"""Raw-text extraction for uploaded documents. Imports of pypdf/python-docx
are deferred into each function so this module can always be imported even
if one of those optional packages isn't installed yet -- the AI Center's
paste-text workflow (and every other page) must never break because of it.
A missing package surfaces as a clear ImportError message from the specific
extractor function, not a crash at import time."""


def extract_text_from_pdf(file_bytes):
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("PDF import requires the 'pypdf' package. Install it with: pip install pypdf")

    import io
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_text_from_docx(file_bytes):
    try:
        import docx
    except ImportError:
        raise ImportError("DOCX import requires the 'python-docx' package. Install it with: pip install python-docx")

    import io
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_csv(file_bytes):
    """Item 2 (Multi-Format Input Support): a spreadsheet of strategy rules
    (one common way strategies actually arrive -- e.g. a row per entry/exit
    rule, or a parameter table) turned into the same '|' -joined row format
    extract_text_from_docx already uses for its tables, so the downstream
    AI prompt sees one consistent tabular text shape regardless of source
    format."""
    import csv
    import io
    text = file_bytes.decode("utf-8-sig", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    lines = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows]
    return "\n".join(line for line in lines if line).strip()


def extract_text_from_xlsx(file_bytes):
    """Item 2 (Multi-Format Input Support): every sheet, in order, each row
    joined the same way as the CSV/DOCX-table extractors above -- one
    consistent tabular text shape for the AI prompt regardless of which
    spreadsheet format the CEO actually pasted in."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("XLSX import requires the 'openpyxl' package. Install it with: pip install openpyxl")

    import io
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(filename, file_bytes):
    """Dispatch by extension. YouTube transcripts and plain notes are
    expected to be pasted directly as text, not uploaded as files, so no
    .txt-specific handling is needed beyond decoding.

    Image-based OCR (a photo/screenshot of a strategy) is deliberately NOT
    supported here -- it needs a real OCR engine (Tesseract) installed on
    the host, which this environment does not have, and adding a wrapper
    around an engine that cannot actually be exercised/tested here would
    claim support that was never proven to work. Left as a known,
    documented gap rather than a guess."""
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return file_bytes.decode("utf-8", errors="replace")
    if lower.endswith(".csv"):
        return extract_text_from_csv(file_bytes)
    if lower.endswith(".xlsx"):
        return extract_text_from_xlsx(file_bytes)
    raise ValueError(f"Unsupported file type for '{filename}'. Supported: .pdf, .docx, .txt, .md, .csv, .xlsx")
